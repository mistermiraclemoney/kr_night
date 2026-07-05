from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_kr_night_business_plan import (
    BLUE,
    CONTENT_WIDTH_DXA,
    CYAN,
    DARK_BLUE,
    FONT_NAME,
    GRID,
    INK,
    LIGHT,
    MUTED,
    NAVY,
    PALE_BLUE,
    PALE_PINK,
    PINK,
    TABLE_INDENT_DXA,
    add_bullet,
    add_callout,
    add_custom_numbering,
    add_heading,
    add_kicker,
    add_page_field,
    add_paragraph,
    add_table,
    rgb,
    set_run_font,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "deliverables" / "KR_NIGHT_사업기획서_2페이지_요약본.docx"


def configure_document(doc: Document) -> None:
    """compact_reference_guide preset with a KR NIGHT Korean-font brand override."""
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.70)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.30)
    section.footer_distance = Inches(0.30)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    heading_tokens = {
        1: (14.5, BLUE, 7, 3),
        2: (12.2, DARK_BLUE, 5, 2),
        3: (11.2, DARK_BLUE, 4, 2),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.08
        style.paragraph_format.keep_with_next = True

    add_custom_numbering(doc)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("KR NIGHT  |  2-PAGE BUSINESS BRIEF"), size=8, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("2026.07  ·  "), size=8, color=MUTED)
    add_page_field(p)


def add_title_block(doc: Document) -> None:
    add_kicker(doc, "BUSINESS PLAN · EXECUTIVE BRIEF", color=PINK, after=2)
    add_paragraph(
        doc,
        "KR NIGHT",
        size=22,
        color=NAVY,
        bold=True,
        after=0,
        line_spacing=1.0,
    )
    add_paragraph(
        doc,
        "Tonight in South Korea",
        size=11.5,
        color=BLUE,
        bold=True,
        after=3,
        line_spacing=1.0,
    )
    add_paragraph(
        doc,
        "한국인과 외국인이 오늘 밤 갈 곳을 발견하고, QR 체크인으로 포인트·현장 라운지·친구 기능을 "
        "이용하는 나이트라이프 플랫폼",
        size=10.2,
        color=INK,
        after=4,
        line_spacing=1.14,
    )
    add_callout(
        doc,
        "핵심 명제",
        "사용자에게는 ‘실패 없는 밤의 선택’을, 매장에는 ‘측정 가능한 실제 방문’을 제공한다.",
        fill=PALE_BLUE,
        accent=CYAN,
    )


def add_page_one(doc: Document) -> None:
    add_title_block(doc)

    add_heading(doc, "1. 기획 의도와 해결 방식", 1)
    add_paragraph(
        doc,
        "클럽·바·라운지 정보는 SNS와 지도에 흩어져 있다. 외국인은 입장 조건·가격·장르 같은 당일 "
        "정보를 확인하기 어렵고, 매장은 홍보가 실제 방문으로 이어졌는지 측정하기 어렵다.",
        size=9.8,
        after=3,
        line_spacing=1.12,
    )
    add_paragraph(
        doc,
        "KR NIGHT는 탐색 → QR 체크인 → 포인트·Lounge → 재방문을 연결한다. 절대적 인기 순위보다 "
        "사용자의 위치·시간·취향에 맞는, 지금 실제로 갈 수 있는 선택지를 정확히 보여준다.",
        size=9.8,
        after=3,
        line_spacing=1.12,
    )

    add_heading(doc, "2. 핵심 타겟과 가치", 1)
    add_table(
        doc,
        ["타겟", "필요", "KR NIGHT의 가치"],
        [
            ("방한 외국인 20~30대", "언어 장벽 없이 실패 없는 밤", "4개 언어, 입장 조건, 지도·QR 패스"),
            ("한국인 20~30대", "오늘 갈 곳을 빠르게 선택", "큐레이션, 친구 상태, 혜택·공유"),
            ("클럽·바·라운지", "신규 고객과 재방문 성과", "QR 방문·쿠폰·재방문 리포트"),
            ("프로모터·관광", "이벤트 노출과 고객 만족", "타겟 배너, 티켓, 다국어 가이드"),
        ],
        [2200, 3000, 4160],
        header_fill=PALE_PINK,
        first_col_bold=True,
        font_size=8.4,
    )

    add_heading(doc, "3. 제품 구조", 1)
    for item in [
        "Home·Hot Now: 추천 Route, 실제 지도, 근처 장소, 사진·영업 상태·입장 조건",
        "QR Check-in: 방문 증명, 포인트 적립, 장소별 Lounge 권한 생성",
        "Social: 친구 검색·상호 승인, 위치 공유 OFF 기본, 장소 단위 상태",
        "Lounge·Game: 체크인 사용자만 입장하는 6시간 채팅과 가벼운 미니게임",
        "Partner Console: 직원 스캔, 쿠폰·이벤트, 체크인·재방문 대시보드",
    ]:
        p = add_bullet(doc, item)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.04
        for run in p.runs:
            set_run_font(run, size=9.3, color=INK)

    add_heading(doc, "4. 벤치마킹과 차별화", 1)
    add_paragraph(
        doc,
        "JP NIGHT의 다국어·혜택 구조, 덤즈의 한국 클럽 정보, CatchTable의 매장 성과 연결을 참고한다. "
        "KR NIGHT는 이를 복제하지 않고 외국인 UX, 검증 QR 방문, 체크인 전용 Lounge, 장소 단위 위치 공유, "
        "매장 재방문 CRM을 하나의 흐름으로 결합한다.",
        size=9.4,
        after=0,
        line_spacing=1.10,
    )


def add_page_two(doc: Document) -> None:
    doc.add_page_break()
    add_kicker(doc, "BUSINESS MODEL & EXECUTION", color=CYAN, after=2)
    add_heading(doc, "5. 비즈니스 모델", 1)
    add_paragraph(
        doc,
        "사용자가 늘수록 매장 입점 가치가 커지고, 좋은 매장·이벤트·혜택이 늘수록 사용자가 증가하는 "
        "양면시장이다. 초기에는 사용자 수보다 ‘KR NIGHT가 실제 손님을 데려왔다’는 증거를 먼저 만든다.",
        size=10.2,
        after=4,
        line_spacing=1.18,
    )
    add_table(
        doc,
        ["수익원", "과금 방식", "도입 기준"],
        [
            ("매장 구독", "정보 관리·통계·쿠폰·Lounge 월 정액", "파일럿 성과 확인 후"),
            ("추천·배너 광고", "지역·날짜·언어 타겟 캠페인비", "초기부터"),
            ("성과형 과금", "유효 QR 체크인·쿠폰 사용 건당", "부정 방지 검증 후"),
            ("거래·VIP", "티켓·예약 수수료와 사용자 프리미엄 구독", "제휴·재방문 확보 후"),
        ],
        [2200, 4400, 2760],
        header_fill=PALE_BLUE,
        first_col_bold=True,
        font_size=8.6,
    )

    add_heading(doc, "6. 시장 진입·마케팅 계획", 1)
    add_table(
        doc,
        ["단계", "핵심 실행", "판단 기준"],
        [
            ("0~30일", "이태원 30곳 정보 구축, 핵심 제휴 10곳, QR 파일럿 3곳", "주간 정보 갱신 80%"),
            ("31~60일", "숏폼·호스텔·대학·외국인 커뮤니티, 현장 QR 유입", "월 활성 1,000명"),
            ("61~90일", "포인트·추천 Route·친구 초대·주간 이벤트", "30일 재방문 20%"),
            ("4~6개월", "홍대 복제, 매장 대시보드, 유료 파일럿", "유료 파트너 10곳"),
        ],
        [1500, 5200, 2660],
        header_fill=PALE_PINK,
        first_col_bold=True,
        font_size=8.7,
    )
    add_paragraph(
        doc,
        "핵심 채널: ‘Tonight in Seoul’ 숏폼, 외국인 입장 가능 장소, 장르별 Route, "
        "입구·테이블 QR, 지역 마이크로 크리에이터, 호스텔·호텔·대학 국제처, 친구 초대 보상.",
        size=9.8,
        color=INK,
        after=3,
        line_spacing=1.13,
    )

    add_heading(doc, "7. 운영 원칙과 실행 기준", 1)
    for item in [
        "초기 운영: 대표 1명 + 현장·제휴 운영 1명, 개발·법률·디자인은 필요한 범위만 외부 활용",
        "연령 확인: 가입 단계의 전면 KYC보다 매장 직원의 실물 신분증 확인을 기본으로 운영",
        "프라이버시: 위치 공유 OFF 기본, 장소 단위 표시, 체크아웃·시간 만료 시 즉시 종료",
        "안전: 신고·차단·관리자 퇴장·QR 재사용 방지·포인트 이상 탐지 체계 필수",
    ]:
        p = add_bullet(doc, item)
        p.paragraph_format.space_after = Pt(1.8)
        p.paragraph_format.line_spacing = 1.05
        for run in p.runs:
            set_run_font(run, size=9.3, color=INK)

    add_callout(
        doc,
        "Go 기준",
        "3개 이상 매장에서 QR 운영이 반복되고, 90일 내 유효 체크인 500건·30일 재방문 20%를 검증하며, "
        "최소 3개 매장이 무료 파일럿 후 유료 또는 성과형 계약 의사를 보이는가.",
        fill=PALE_BLUE,
        accent=CYAN,
    )
    add_paragraph(
        doc,
        "다음 행동: 이태원 제휴 3곳에서 4주간 ‘정보 → 체크인 → 혜택 → 재방문’ 전체 흐름을 실제 운영한다. "
        "가격·일정·KPI는 검증 가설이며, 위치정보·개인정보·연령 확인은 출시 전 전문 검토가 필요하다.",
        size=8.8,
        color=MUTED,
        italic=True,
        after=0,
        line_spacing=1.10,
    )


def audit_tables(doc: Document) -> None:
    for index, table in enumerate(doc.tables, start=1):
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
        if tbl_w is None or tbl_w.get(qn("w:w")) != str(CONTENT_WIDTH_DXA):
            raise AssertionError(f"Table {index}: bad tblW")
        if tbl_ind is None or tbl_ind.get(qn("w:w")) != str(TABLE_INDENT_DXA):
            raise AssertionError(f"Table {index}: bad tblInd")
        grid_widths = [int(node.get(qn("w:w"))) for node in table._tbl.tblGrid]
        if sum(grid_widths) != CONTENT_WIDTH_DXA:
            raise AssertionError(f"Table {index}: bad grid width sum {sum(grid_widths)}")
        for row in table.rows:
            for cell_index, cell in enumerate(row.cells):
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                if tc_w is None or int(tc_w.get(qn("w:w"))) != grid_widths[cell_index]:
                    raise AssertionError(f"Table {index}: bad cell width")


def build() -> Path:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_page_one(doc)
    add_page_two(doc)
    audit_tables(doc)

    core = doc.core_properties
    core.title = "KR NIGHT 사업기획서 2페이지 요약본"
    core.subject = "나이트라이프 발견·QR 체크인·소셜 플랫폼 사업 요약"
    core.author = "KR NIGHT"
    core.keywords = "KR NIGHT, nightlife, QR check-in, business brief"
    core.comments = "2-page executive brief, 2026-07"
    doc.save(OUTPUT)
    print(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    build()
