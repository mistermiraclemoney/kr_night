from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
QA_DIR = ROOT / "qa" / "kr-night-business-plan"
ASSET_DIR = ROOT / "assets"
OUTPUT = OUT_DIR / "KR_NIGHT_사업기획서_v1.0.docx"
FLYWHEEL_IMAGE = QA_DIR / "kr-night-flywheel.png"

FONT_NAME = "Apple SD Gothic Neo"
FONT_FILE = "/System/Library/Fonts/AppleSDGothicNeo.ttc"

# narrative_proposal preset, with a named KR NIGHT brand override.
INK = "101522"
NAVY = "17233A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
CYAN = "00BFE8"
PINK = "F01A8B"
MUTED = "657085"
LIGHT = "F4F6F9"
PALE_BLUE = "E9F7FC"
PALE_PINK = "FCEBF4"
WHITE = "FFFFFF"
GRID = "D8DEE8"
GOLD = "C79326"
GREEN = "16815D"
RED = "A33A4A"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(
    run,
    *,
    name: str = FONT_NAME,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=GRID, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_hyperlink(paragraph, text: str, url: str, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        run_fonts.set(qn(f"w:{attr}"), FONT_NAME)
    run_pr.append(run_fonts)
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    run_pr.append(run_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.append(underline)
    run.append(run_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, display, end])
    set_run_font(run, size=9, color=MUTED)


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    style: str | None = None,
    size: float = 11,
    color: str = INK,
    bold: bool = False,
    italic: bool = False,
    align=None,
    before: float = 0,
    after: float = 8,
    line_spacing: float = 1.333,
    keep_next: bool = False,
):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    if keep_next:
        set_keep_with_next(paragraph)
    return paragraph


def add_mixed_paragraph(doc: Document, chunks: Sequence[tuple[str, bool]], after=8):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.333
    for text, bold in chunks:
        set_run_font(paragraph.add_run(text), size=11, color=INK, bold=bold)
    return paragraph


def add_bullet(doc: Document, text: str, level: int = 0):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.194)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "899")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    set_run_font(paragraph.add_run(text), size=11, color=INK)
    return paragraph


def add_numbered(doc: Document, text: str, level: int = 0, num_id_value: int = 900):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.194)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(num_id_value))
    num_pr.append(ilvl)
    num_pr.append(num_id)
    set_run_font(paragraph.add_run(text), size=11, color=INK)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    set_keep_with_next(paragraph)
    run = paragraph.add_run(text)
    return paragraph


def add_kicker(doc: Document, text: str, *, color=CYAN, after=4, align=WD_ALIGN_PARAGRAPH.LEFT):
    return add_paragraph(
        doc,
        text.upper(),
        size=9.5,
        color=color,
        bold=True,
        align=align,
        after=after,
        line_spacing=1.0,
        keep_next=True,
    )


def add_callout(doc: Document, label: str, text: str, fill=PALE_BLUE, accent=CYAN):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=accent, size=8)
    set_cell_shading(table.cell(0, 0), fill)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run_font(paragraph.add_run(f"{label}  "), size=10.5, color=accent, bold=True)
    set_run_font(paragraph.add_run(text), size=10.5, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return table


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths_dxa: Sequence[int],
    *,
    header_fill=LIGHT,
    first_col_bold=False,
    font_size=9.4,
):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        set_run_font(p.add_run(header), size=9.2, color=NAVY, bold=True)

    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.18
            if i > 0 and len(str(value)) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(
                p.add_run(str(value)),
                size=font_size,
                color=INK,
                bold=first_col_bold and i == 0,
            )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(5)
    return table


def add_source_note(doc: Document, text: str):
    return add_paragraph(
        doc,
        text,
        size=8.5,
        color=MUTED,
        italic=True,
        before=4,
        after=4,
        line_spacing=1.15,
    )


def add_horizontal_rule(paragraph, color=GRID, size=8):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_custom_numbering(doc: Document):
    numbering = doc.part.numbering_part.element

    def add_abstract(abstract_id: int, fmt: str, text: str, font: str | None = None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "multilevel")
        abstract.append(multi)
        for level in range(3):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(level))
            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            lvl.append(start)
            num_fmt = OxmlElement("w:numFmt")
            num_fmt.set(qn("w:val"), fmt)
            lvl.append(num_fmt)
            lvl_text = OxmlElement("w:lvlText")
            lvl_text.set(qn("w:val"), text if fmt == "bullet" else f"%{level + 1}.")
            lvl.append(lvl_text)
            jc = OxmlElement("w:lvlJc")
            jc.set(qn("w:val"), "left")
            lvl.append(jc)
            p_pr = OxmlElement("w:pPr")
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "num")
            tab.set(qn("w:pos"), str(540 + level * 360))
            tabs.append(tab)
            p_pr.append(tabs)
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(540 + level * 360))
            ind.set(qn("w:hanging"), "280")
            p_pr.append(ind)
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:after"), "80")
            spacing.set(qn("w:line"), "290")
            spacing.set(qn("w:lineRule"), "auto")
            p_pr.append(spacing)
            lvl.append(p_pr)
            if font:
                r_pr = OxmlElement("w:rPr")
                r_fonts = OxmlElement("w:rFonts")
                r_fonts.set(qn("w:ascii"), font)
                r_fonts.set(qn("w:hAnsi"), font)
                r_pr.append(r_fonts)
                lvl.append(r_pr)
            abstract.append(lvl)
        numbering.append(abstract)

    def add_num(num_id: int, abstract_id: int, restart_at_one: bool = False):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(abstract_id))
        num.append(abstract_num_id)
        if restart_at_one:
            override = OxmlElement("w:lvlOverride")
            override.set(qn("w:ilvl"), "0")
            start_override = OxmlElement("w:startOverride")
            start_override.set(qn("w:val"), "1")
            override.append(start_override)
            num.append(override)
        numbering.append(num)

    add_abstract(899, "bullet", "•", "Arial")
    add_num(899, 899)
    add_abstract(900, "decimal", "%1.")
    add_num(900, 900)
    add_num(901, 900, restart_at_one=True)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_tokens = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 12, 6),
        3: (12, DARK_BLUE, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.12
        style.paragraph_format.keep_with_next = True

    add_custom_numbering(doc)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    set_run_font(header_p.add_run("KR NIGHT  |  BUSINESS PLAN"), size=8.5, color=MUTED, bold=True)
    add_horizontal_rule(header_p, color=GRID, size=6)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    set_run_font(footer_p.add_run("KR NIGHT · 2026.07 · "), size=8.5, color=MUTED)
    add_page_field(footer_p)


def make_flywheel():
    QA_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 1700, 430
    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    font_bold = ImageFont.truetype(FONT_FILE, 37, index=5)
    font_small = ImageFont.truetype(FONT_FILE, 24, index=0)

    boxes = [
        ("매장·이벤트 정보", "좋은 밤의 선택지"),
        ("사용자 발견", "검색·추천·지도"),
        ("QR 체크인", "방문을 데이터로 증명"),
        ("포인트·소셜", "현장 경험과 재방문"),
        ("성과·유료 제휴", "매장 ROI와 공급 확대"),
    ]
    colors = [CYAN, "4C8DE8", "7A64E8", PINK, "D75B75"]
    box_w, box_h, gap = 285, 180, 36
    start_x = 45
    y = 105
    for index, ((title, subtitle), color) in enumerate(zip(boxes, colors)):
        x = start_x + index * (box_w + gap)
        draw.rounded_rectangle(
            (x, y, x + box_w, y + box_h),
            radius=26,
            fill="#F7F9FC",
            outline=f"#{color}",
            width=5,
        )
        title_bbox = draw.textbbox((0, 0), title, font=font_bold)
        title_x = x + (box_w - (title_bbox[2] - title_bbox[0])) / 2
        draw.text((title_x, y + 42), title, font=font_bold, fill=f"#{INK}")
        subtitle_bbox = draw.textbbox((0, 0), subtitle, font=font_small)
        subtitle_x = x + (box_w - (subtitle_bbox[2] - subtitle_bbox[0])) / 2
        draw.text((subtitle_x, y + 111), subtitle, font=font_small, fill=f"#{MUTED}")
        if index < len(boxes) - 1:
            arrow_x = x + box_w + 7
            arrow_y = y + box_h / 2
            draw.line((arrow_x, arrow_y, arrow_x + gap - 15, arrow_y), fill=f"#{MUTED}", width=4)
            draw.polygon(
                [
                    (arrow_x + gap - 15, arrow_y),
                    (arrow_x + gap - 28, arrow_y - 9),
                    (arrow_x + gap - 28, arrow_y + 9),
                ],
                fill=f"#{MUTED}",
            )
    draw.arc((70, 18, width - 70, height - 8), start=190, end=350, fill=f"#{PINK}", width=5)
    draw.polygon([(68, 224), (88, 207), (91, 231)], fill=f"#{PINK}")
    image.save(FLYWHEEL_IMAGE)


def add_cover(doc: Document):
    add_paragraph(doc, "", after=26)
    logo = ASSET_DIR / "kr-night-app-icon.png"
    picture_p = doc.add_paragraph()
    picture_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_p.paragraph_format.space_after = Pt(18)
    picture_p.add_run().add_picture(str(logo), width=Inches(2.35))

    add_kicker(doc, "BUSINESS PLAN · VERSION 1.0", color=PINK, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        doc,
        "KR NIGHT",
        size=30,
        color=NAVY,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=4,
        line_spacing=1.0,
    )
    add_paragraph(
        doc,
        "Tonight in South Korea",
        size=15,
        color=BLUE,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=18,
        line_spacing=1.0,
    )
    add_paragraph(
        doc,
        "한국인과 외국인을 연결하는\n나이트라이프 발견·체크인·소셜 플랫폼",
        size=14,
        color=INK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=34,
        line_spacing=1.4,
    )

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(18)
    add_horizontal_rule(rule, color=CYAN, size=14)
    add_paragraph(
        doc,
        "작성일  2026년 7월 3일\n문서 목적  사업 구조 정리 · 제휴 설명 · 실행 기준 수립",
        size=10,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=6,
        line_spacing=1.35,
    )
    add_paragraph(
        doc,
        "본 문서의 가격·매출·일정 수치는 시장 검증을 위한 가설이며, 확정 수치가 아닙니다.",
        size=8.5,
        color=MUTED,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=0,
        line_spacing=1.2,
    )
    doc.add_page_break()


def add_document_map(doc: Document):
    add_kicker(doc, "DOCUMENT MAP", color=CYAN)
    add_heading(doc, "문서 구성", 1)
    add_paragraph(
        doc,
        "이 기획서는 KR NIGHT가 무엇인지, 누구의 문제를 해결하는지, 어떻게 사용자를 모으고 "
        "매장에 성과를 제공하며 수익으로 전환하는지를 하나의 사업 구조로 설명한다.",
    )
    sections = [
        ("01", "Executive Summary", "사업의 핵심 명제와 우선순위"),
        ("02", "기획 의도와 문제 정의", "사용자·매장·시장에 존재하는 불편"),
        ("03", "타겟과 고객 경험", "핵심 사용자, 매장, 프로모터"),
        ("04", "제품 및 서비스 구조", "발견, QR, 포인트, 소셜, 게임"),
        ("05", "시장·벤치마킹·차별화", "JP NIGHT, 덤즈, CatchTable 등"),
        ("06", "비즈니스·수익 모델", "B2B, 거래 수수료, 광고, 멤버십"),
        ("07", "마케팅·성장 계획", "공급 확보, 바이럴, 재방문"),
        ("08", "운영·기술·안전", "현실적인 운영 체계와 준법 과제"),
        ("09", "재무 가설·KPI·로드맵", "측정 가능한 실행 기준"),
        ("10", "제휴 제안과 결론", "매장에 전달할 메시지"),
    ]
    add_table(
        doc,
        ["구분", "섹션", "핵심 질문"],
        sections,
        [800, 3000, 5560],
        header_fill=PALE_BLUE,
        first_col_bold=True,
        font_size=9.8,
    )
    add_callout(
        doc,
        "한 문장 정의",
        "KR NIGHT는 오늘 밤 갈 곳을 찾는 사람과 손님을 원하는 매장을 연결하고, "
        "QR 체크인으로 실제 방문을 증명해 포인트·소셜·성과 데이터를 만드는 나이트라이프 플랫폼이다.",
        fill=PALE_PINK,
        accent=PINK,
    )


def section_executive_summary(doc: Document):
    add_kicker(doc, "01 · EXECUTIVE SUMMARY", color=PINK)
    add_heading(doc, "사업 요약", 1)
    add_paragraph(
        doc,
        "한국의 나이트라이프 정보는 인스타그램, 틱톡, 지도 리뷰, 지인 추천, 개별 매장 계정에 "
        "흩어져 있다. 특히 외국인 방문객은 입장 조건, 외국인 입장 가능 여부, 음악 장르, "
        "당일 라인업, 가격, 드레스코드 같은 ‘방문 직전 정보’를 한 번에 확인하기 어렵다. "
        "매장 역시 SNS 홍보만으로는 실제 방문 전환을 측정하기 어렵다.",
    )
    add_paragraph(
        doc,
        "KR NIGHT는 이 문제를 단순한 장소 목록이 아니라 하나의 행동 흐름으로 해결한다. "
        "사용자는 장소를 발견하고, 매장에서 QR로 체크인하며, 포인트와 현장 라운지를 이용하고, "
        "친구와 밤의 경험을 공유한다. 매장은 노출을 넘어 실제 체크인·쿠폰 사용·재방문 데이터를 받는다.",
    )

    add_table(
        doc,
        ["항목", "정의"],
        [
            ("비전", "한국의 밤을 가장 쉽게 발견하고, 안전하게 연결하고, 다시 찾게 만드는 플랫폼"),
            ("핵심 고객", "20~30대 한국인, 방한 외국인, 클럽·바·라운지, 이벤트 프로모터"),
            ("핵심 행동", "탐색 → 장소 선택 → QR 체크인 → 포인트·라운지 → 재방문·추천"),
            ("초기 지역", "이태원 중심 검증 후 홍대, 압구정·신사, 강남, 성수, 을지로로 확장"),
            ("초기 수익", "유료 노출, 매장 구독, 캠페인비, 티켓·예약·검증 방문 성과 수수료"),
            ("장기 자산", "검증된 방문 데이터, 사용자 취향 그래프, 매장 성과 데이터, 파트너 네트워크"),
        ],
        [1900, 7460],
        header_fill=LIGHT,
        first_col_bold=True,
        font_size=9.7,
    )
    add_callout(
        doc,
        "가장 중요한 원칙",
        "처음부터 모든 기능을 완성하는 것이 목표가 아니다. 사용자가 ‘오늘 어디 가지?’에 답하고, "
        "매장이 ‘KR NIGHT가 실제 손님을 데려왔다’를 확인하는 순간을 먼저 완성한다.",
        fill=PALE_BLUE,
        accent=CYAN,
    )


def section_problem_intent(doc: Document):
    add_kicker(doc, "02 · INTENT & PROBLEM", color=CYAN)
    add_heading(doc, "기획 의도와 문제 정의", 1)

    add_heading(doc, "2.1 왜 지금 KR NIGHT인가", 2)
    add_paragraph(
        doc,
        "2026년 1~4월 서울 외국인 방문객은 520만 명으로 전년 동기 대비 21.4% 증가했고, "
        "2026년 4월 외국인 방문객의 서울 카드 소비는 약 1.153조 원이었다. 서울시는 소비가 "
        "쇼핑뿐 아니라 식음료·체험형 영역으로 확장되고 있으며, 홍대·성수 등 신흥 상권으로도 "
        "퍼지고 있다고 분석했다.[1] 이는 다국어로 지역 경험을 연결하는 플랫폼에 유리한 환경이다.",
    )
    add_source_note(doc, "근거: 서울특별시 2026년 6월 발표 및 한국관광공사 관광데이터랩 인용.[1]")

    add_heading(doc, "2.2 사용자 문제", 2)
    for item in [
        "당일 라인업, 입장료, 피크 시간, 음악 장르, 드레스코드가 여러 채널에 흩어져 있다.",
        "외국인은 여권 필요 여부, 외국인 입장 가능 여부, 언어 지원 등 추가 정보를 필요로 한다.",
        "지도 리뷰는 낮 시간대 정보와 섞여 있어 ‘오늘 밤의 상태’를 판단하기 어렵다.",
        "낯선 장소에서는 함께 갈 사람, 현장 분위기, 안전한 소통 수단이 부족하다.",
        "쿠폰과 혜택이 있어도 어디서 어떻게 쓰는지 명확하지 않으면 방문으로 이어지지 않는다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.3 매장 문제", 2)
    for item in [
        "인스타그램 게시물과 광고가 실제 입장 고객으로 전환됐는지 측정하기 어렵다.",
        "평일·이른 시간·특정 이벤트의 빈 좌석을 채울 정밀한 유입 채널이 부족하다.",
        "외국인 고객에게 입장 규칙과 혜택을 반복 설명하는 운영 비용이 크다.",
        "프로모터·게스트리스트·쿠폰 성과가 수기로 관리돼 데이터가 매장 자산으로 남지 않는다.",
        "여러 플랫폼에 정보를 중복 업데이트해야 하며, 당일 변경 사항 반영이 늦다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.4 기획의 핵심 전환", 2)
    add_mixed_paragraph(
        doc,
        [
            ("KR NIGHT가 해결해야 하는 것은 ", False),
            ("‘어디가 제일 핫한가’라는 모호한 인기 순위가 아니다. ", True),
            ("사용자의 위치·취향·시간에 맞는 선택지를 보여주고, 실제 방문과 혜택 사용을 증명하는 것이다.", False),
        ],
    )
    add_callout(
        doc,
        "제품 명제",
        "Hot Now는 절대적인 혼잡도 순위보다 ‘지금 갈 수 있는 내 주변 선택지’를 정확하고 보기 좋게 "
        "제공해야 한다. 혼잡도는 데이터가 쌓인 뒤 보조 신호로 추가한다.",
        fill=PALE_PINK,
        accent=PINK,
    )


def section_targets(doc: Document):
    add_kicker(doc, "03 · TARGET & EXPERIENCE", color=PINK)
    add_heading(doc, "타겟과 고객 경험", 1)

    add_heading(doc, "3.1 핵심 타겟", 2)
    add_table(
        doc,
        ["타겟", "핵심 욕구", "KR NIGHT의 가치"],
        [
            ("방한 외국인 20~30대", "언어 장벽 없이 실패 없는 밤을 보내고 싶다", "다국어 정보, 외국인 입장 조건, 지도, QR 패스"),
            ("한국인 20~30대", "오늘 갈 곳과 이벤트를 빠르게 고르고 싶다", "큐레이션, 친구 상태, 혜택, 저장·공유"),
            ("클럽·바·라운지", "실제 방문 고객과 재방문을 만들고 싶다", "QR 성과 측정, 쿠폰, 프로모션, 매장 대시보드"),
            ("DJ·래퍼·프로모터", "라인업과 티켓을 타겟 고객에게 노출하고 싶다", "이벤트 배너, 티켓 연결, 팬 유입 데이터"),
            ("호텔·호스텔·관광 파트너", "투숙객에게 신뢰할 수 있는 밤 정보를 주고 싶다", "지역별 다국어 가이드와 제휴 코드"),
        ],
        [1900, 3100, 4360],
        header_fill=PALE_BLUE,
        first_col_bold=True,
        font_size=9.2,
    )

    add_heading(doc, "3.2 우선 공략 순서", 2)
    add_numbered(doc, "이태원: 외국인 친화성과 다국적 수요가 강해 다국어 가치 검증이 빠르다.")
    add_numbered(doc, "홍대: 대학생·여행객·힙합·라이브 클럽이 밀집해 콘텐츠와 바이럴이 강하다.")
    add_numbered(doc, "압구정·신사·강남: 객단가와 VIP 테이블 수요가 높아 거래 수익 모델을 검증하기 좋다.")
    add_numbered(doc, "성수·을지로: 바·팝업·문화 이벤트 중심으로 ‘클럽 외 밤문화’ 범위를 확장한다.")

    add_heading(doc, "3.3 대표 사용자 여정", 2)
    add_table(
        doc,
        ["단계", "사용자 행동", "제품 접점", "성공 신호"],
        [
            ("발견", "오늘 밤 갈 곳을 탐색", "Home, Hot Now, 검색, 추천 Route", "장소 상세 조회"),
            ("판단", "가격·장르·입장 조건 비교", "사진, 이벤트, 다국어 규칙, 후기", "저장·공유·길찾기"),
            ("입장", "매장 직원에게 QR 제시", "회전형 QR Pass, 직원 스캐너", "검증 체크인"),
            ("현장", "포인트·라운지·친구 이용", "Lounge, Friends, Game", "메시지·친구 추가"),
            ("재방문", "포인트 사용·다음 일정 선택", "쿠폰, 알림, 추천 Route", "30일 재방문"),
        ],
        [1200, 2500, 3100, 2560],
        header_fill=LIGHT,
        first_col_bold=True,
        font_size=8.9,
    )


def section_product(doc: Document):
    add_kicker(doc, "04 · PRODUCT", color=CYAN)
    add_heading(doc, "제품 및 서비스 구조", 1)

    add_heading(doc, "4.1 다섯 개의 사용자 탭", 2)
    add_table(
        doc,
        ["탭", "역할", "핵심 기능"],
        [
            ("Home", "밤의 시작점", "개인화 추천, 오늘의 Route, 이벤트 배너, 지역·언어 전환"),
            ("Social", "친구와 현장 연결", "Friends, 메시지, 선택형 위치 공유, 장소별 Lounge"),
            ("Hot Now", "근처 장소 탐색", "실제 지도, 사진 카드, 필터, 거리·영업 상태, 길찾기"),
            ("Game", "현장 체류와 상호작용", "아이스브레이커, 술래 추첨, 그룹 미니게임, 추후 매칭"),
            ("Profile", "멤버십과 신뢰", "로그인, 닉네임·ID, QR Pass, 포인트, 쿠폰, 언어, 개인정보 설정"),
        ],
        [1200, 2100, 6060],
        header_fill=PALE_PINK,
        first_col_bold=True,
        font_size=9.2,
    )

    add_heading(doc, "4.2 핵심 엔진: QR 체크인", 2)
    add_paragraph(
        doc,
        "QR은 단순 입장권이 아니라 KR NIGHT 사업모델의 중심 데이터다. 사용자의 방문을 증명하고, "
        "포인트 적립과 장소 라운지 권한을 열며, 매장에는 캠페인 성과를 제공한다.",
    )
    for item in [
        "사용자 QR은 고정 이미지가 아니라 짧은 시간마다 바뀌는 서명 토큰으로 운영한다.",
        "직원 앱 또는 매장 대시보드가 QR을 스캔하고, 필요 시 실물 신분증과 표시 이름을 대조한다.",
        "체크인 성공 시 서버가 방문 기록, 포인트, 라운지 권한, 선택형 장소 상태를 동시에 생성한다.",
        "중복 스캔, 캡처 이미지 재사용, 비정상 포인트 적립을 탐지한다.",
        "매장 퇴장·영업 종료·일정 시간 경과 시 라운지 권한과 위치 상태가 자동 만료된다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.3 현장 Lounge의 운영 원칙", 2)
    add_paragraph(
        doc,
        "Lounge는 체크인한 사람만 들어오는 장소 단위 오픈채팅이다. 매장 내 대화를 촉진하되, "
        "상시 공개 커뮤니티가 되지 않도록 권한을 시간 제한형으로 설계한다.",
    )
    add_table(
        doc,
        ["항목", "권장 정책"],
        [
            ("입장", "해당 매장의 유효한 QR 체크인 보유자만"),
            ("만료", "체크인 후 6시간, 매장 영업 종료, 명시적 체크아웃 중 가장 빠른 시점"),
            ("위치", "정확한 좌표가 아닌 ‘Henz Club에 있음’ 같은 장소 단위 상태"),
            ("메시지", "24시간 후 자동 삭제를 기본으로 하되 신고 증거는 별도 보존 정책 적용"),
            ("안전", "차단·신고·금칙어·관리자 퇴장·재입장 제한"),
            ("친구 전환", "상호 동의한 경우에만 Lounge 관계가 Friends로 이어짐"),
        ],
        [1800, 7560],
        header_fill=LIGHT,
        first_col_bold=True,
        font_size=9.6,
    )

    add_heading(doc, "4.4 포인트와 리워드", 2)
    add_paragraph(
        doc,
        "포인트는 무조건 많이 주는 화폐가 아니라 행동 설계 도구다. 방문, 추천, 콘텐츠 기여처럼 "
        "사업에 가치가 생기는 행동에만 지급하고, 매장 혜택과 교환되도록 한다.",
    )
    for item in [
        "적립: 첫 체크인, 비혼잡 시간 방문, 친구 초대 후 첫 방문, 검증된 후기, 이벤트 참여",
        "사용: 맥주·웰컴드링크 쿠폰, 입장료 할인, 게스트리스트, 한정 이벤트 응모",
        "재원: 제휴 매장 공동 부담, 캠페인 예산, KR NIGHT 마케팅 예산",
        "방지: 일일 적립 한도, 동일 매장 재적립 제한, 기기·계정 이상 탐지",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.5 게임과 매칭의 위치", 2)
    add_paragraph(
        doc,
        "게임은 앱의 중심이 아니라 현장 체류와 친구 형성을 돕는 보조 기능이다. 미니게임은 먼저 "
        "운영할 수 있지만, 랜덤 매칭·VIP 이상형 필터는 연령 확인, 신고 대응, 성비 불균형, "
        "괴롭힘 방지 체계가 준비된 뒤 제한된 매장에서 검증해야 한다.",
    )


def section_market_benchmark(doc: Document):
    add_kicker(doc, "05 · MARKET & BENCHMARK", color=PINK)
    add_heading(doc, "시장, 벤치마킹 사례, 차별화", 1)

    add_heading(doc, "5.1 시장 기회", 2)
    add_paragraph(
        doc,
        "서울의 외국인 방문과 체험 소비가 성장하고, 이태원·홍대·강남·성수처럼 서로 다른 밤 문화를 "
        "가진 상권이 공존한다. 그러나 정보·예약·입장·소셜이 하나의 사용자 계정과 방문 데이터로 "
        "연결된 플랫폼은 아직 분절돼 있다. KR NIGHT의 기회는 ‘정보가 없다’가 아니라 "
        "‘정보 이후의 방문과 재방문이 연결되지 않는다’는 데 있다.",
    )

    add_heading(doc, "5.2 벤치마킹 사례", 2)
    add_table(
        doc,
        ["서비스", "배울 점", "한계 또는 빈 공간", "KR NIGHT 적용"],
        [
            (
                "JP NIGHT",
                "점포 검색, 할인, 티켓, 포인트, 4개 언어, 제휴 생태계",
                "일본 시장 중심이며 한국 현장 운영 데이터는 없음",
                "다국어·혜택 구조를 참고하되 QR 방문·현장 소셜로 확장",
            ),
            (
                "덤즈",
                "한국 클럽·라운지 정보, 쿠폰, 실시간 입장료, 공연·파티, 매장 홍보",
                "KR NIGHT와 직접 경쟁하는 정보·쿠폰 영역",
                "단순 목록 경쟁을 피하고 검증 방문·외국인 UX·CRM 성과에 집중",
            ),
            (
                "CatchTable",
                "발견과 예약·대기·매장 CRM을 연결하고 검증 방문 기반 리뷰를 구축",
                "외식 중심이며 밤문화의 입장·소셜 맥락은 다름",
                "소비자 앱과 매장 운영 도구를 함께 키우는 양면시장 모델",
            ),
            (
                "Resident Advisor",
                "장르 기반 이벤트 발견, 아티스트·클럽 디렉터리, 티켓 판매",
                "전자음악 중심, 한국 로컬 입장 운영과 다국어 장벽",
                "DJ·이벤트 정보 품질과 티켓 거래 구조 참고",
            ),
            (
                "Snap Map",
                "친구 기반 선택형 위치 공유와 지도형 소셜 경험",
                "정확한 위치 노출에 따른 안전·프라이버시 위험",
                "위치 공유 OFF 기본, 상호 친구·장소 단위·자동 만료 적용",
            ),
            (
                "NightFlow",
                "외국인을 위한 서울 VIP 테이블 예약과 영어 지원",
                "VIP 예약이라는 좁은 거래 중심",
                "일반 방문부터 VIP까지 이어지는 상위 전환 경로로 흡수",
            ),
        ],
        [1300, 2450, 2600, 3010],
        header_fill=PALE_BLUE,
        first_col_bold=True,
        font_size=8.45,
    )
    add_source_note(doc, "벤치마킹 근거: 각 서비스 공식 사이트 및 공식 지원 문서.[2][3][4][5][6][7]")

    add_heading(doc, "5.3 경쟁 현실", 2)
    add_callout(
        doc,
        "정직한 결론",
        "KR NIGHT는 ‘한국 최초의 클럽 정보앱’으로 포지셔닝하면 안 된다. 이미 덤즈 등 유사 서비스가 "
        "존재한다. 승부처는 정보량이 아니라 외국인까지 통하는 사용성, 실제 QR 방문 데이터, "
        "장소 라운지, 재방문 리워드, 매장 성과 도구의 결합이다.",
        fill=PALE_PINK,
        accent=PINK,
    )

    add_heading(doc, "5.4 차별화 포지셔닝", 2)
    add_table(
        doc,
        ["차별화 축", "KR NIGHT의 기준"],
        [
            ("다국어", "한국어·영어·일본어·중국어 전체 UI와 장소별 입장 규칙 번역"),
            ("검증 방문", "QR 체크인으로 방문·쿠폰·포인트·리뷰를 하나의 이벤트로 기록"),
            ("현장 소셜", "체크인한 장소에서만 열리고 자동 만료되는 Lounge"),
            ("프라이버시", "위치 공유 OFF 기본, 장소 단위 표시, 선택 친구, 만료 시간"),
            ("매장 성과", "노출이 아니라 체크인·쿠폰·재방문·시간대별 전환을 리포트"),
            ("콘텐츠", "‘오늘의 추천 Route’와 현장성 있는 다국어 큐레이션"),
        ],
        [1900, 7460],
        header_fill=LIGHT,
        first_col_bold=True,
        font_size=9.6,
    )

    add_heading(doc, "5.5 브랜드와 상표", 2)
    add_paragraph(
        doc,
        "‘KR NIGHT’는 의미가 직관적이지만 ‘JP NIGHT’와 이름 구조가 유사하다. 로고·캐릭터·색상·문구는 "
        "독립된 정체성을 유지하고, 정식 출시 전 한국 및 주요 진출국의 상표 선행조사와 출원 가능성 검토가 필요하다. "
        "이 문서는 상표 등록 가능성을 보증하지 않는다.",
    )


def section_business_model(doc: Document):
    add_kicker(doc, "06 · BUSINESS MODEL", color=CYAN)
    add_heading(doc, "전반적인 비즈니스 모델", 1)

    add_heading(doc, "6.1 양면시장 구조", 2)
    add_paragraph(
        doc,
        "KR NIGHT는 사용자가 많아질수록 매장 입점 가치가 커지고, 좋은 매장·이벤트·혜택이 많아질수록 "
        "사용자가 늘어나는 양면시장이다. 다만 초기에는 사용자 수보다 ‘검증 가능한 매장 성과’를 먼저 만들어야 한다.",
    )
    make_flywheel()
    image_p = doc.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.space_after = Pt(4)
    image_p.add_run().add_picture(str(FLYWHEEL_IMAGE), width=Inches(6.45))
    add_source_note(doc, "KR NIGHT 성장 플라이휠: 공급 → 발견 → 체크인 → 재방문 → 성과 → 공급 확대.")

    add_heading(doc, "6.2 고객별 가치 제안", 2)
    add_table(
        doc,
        ["고객", "얻는 가치", "제공하는 자산"],
        [
            ("사용자", "더 빠른 선택, 혜택, 친구·현장 경험", "관심, 체크인, 후기, 추천"),
            ("매장", "신규 고객, 다국어 안내, 재방문, 성과 측정", "장소 정보, 혜택, 제휴비"),
            ("프로모터", "이벤트 노출, 티켓 전환, 관객 데이터", "라인업, 콘텐츠, 수수료"),
            ("관광 파트너", "고객 만족과 지역 경험 차별화", "추천 채널, 공동 프로모션"),
        ],
        [1600, 3800, 3960],
        header_fill=PALE_PINK,
        first_col_bold=True,
        font_size=9.5,
    )

    add_heading(doc, "6.3 수익 모델", 2)
    add_table(
        doc,
        ["수익원", "과금 대상", "과금 방식", "도입 시점"],
        [
            ("매장 구독", "클럽·바·라운지", "월 정액: 정보 관리, 통계, 쿠폰, 라운지 관리", "초기 성과 입증 후"),
            ("추천·배너 광고", "매장·주류·이벤트", "지역·날짜·언어 타겟 캠페인비", "초기부터 가능"),
            ("검증 방문 성과비", "제휴 매장", "유효 QR 체크인 또는 쿠폰 사용 건당", "부정 방지 후"),
            ("티켓 수수료", "이벤트 주최자", "결제액의 일정 비율 + 결제대행비", "거래 기능 도입 후"),
            ("VIP 테이블·예약", "매장·프로모터", "예약금 또는 매출의 일정 비율", "프리미엄 지역 확장 시"),
            ("사용자 VIP", "헤비 유저", "월 구독: 우선 혜택, 고급 필터, 한정 이벤트", "충분한 혜택 확보 후"),
            ("데이터 리포트", "매장·브랜드", "익명 집계 리포트·캠페인 분석", "데이터 규모 확보 후"),
        ],
        [1750, 1700, 3900, 2010],
        header_fill=LIGHT,
        first_col_bold=True,
        font_size=8.8,
    )

    add_heading(doc, "6.4 가격 가설", 2)
    add_paragraph(
        doc,
        "아래 가격은 영업 현장에서 반응을 확인하기 위한 출발점이다. 처음 10~20개 매장은 무료 또는 "
        "성과형으로 운영하고, 유료 전환 의향을 인터뷰로 검증한다.",
    )
    add_table(
        doc,
        ["플랜", "월 가격 가설", "제공 가치", "대상"],
        [
            ("Free", "0원", "기본 정보, 이벤트 제출, 외부 지도 연결", "초기 공급 확보"),
            ("Growth", "99,000~149,000원", "쿠폰, 기본 통계, 다국어 정보 관리", "일반 바·클럽"),
            ("Pro", "249,000~399,000원", "추천 노출, CRM, 라운지 관리, 고급 통계", "핵심 제휴 매장"),
            ("Campaign", "건별 30만~300만원", "배너·Route·콘텐츠·현장 QR 공동 프로모션", "이벤트·주류 브랜드"),
            ("Performance", "체크인/거래별 협의", "검증 방문, 티켓, 예약 성과 기반", "성과형 선호 파트너"),
        ],
        [1450, 1900, 4010, 2000],
        header_fill=PALE_BLUE,
        first_col_bold=True,
        font_size=9.1,
    )
    add_callout(
        doc,
        "수익화 원칙",
        "사장님이 돈을 내는 이유는 앱이 예뻐서가 아니라 ‘손님이 왔고 다시 왔다’는 데이터가 있기 때문이다. "
        "유료화보다 먼저 10개 매장에서 QR 유입과 재방문 사례를 만든다.",
        fill=PALE_PINK,
        accent=PINK,
    )


def section_marketing(doc: Document):
    add_kicker(doc, "07 · GO-TO-MARKET", color=PINK)
    add_heading(doc, "마케팅 및 성장 계획", 1)

    add_heading(doc, "7.1 시장 진입 전략", 2)
    add_paragraph(
        doc,
        "처음부터 서울 전체를 채우면 정보 품질과 운영 집중도가 무너진다. 한 지역에서 사용자가 "
        "‘오늘 밤 켤 이유’를 느낄 만큼 밀도를 만들고, 검증된 운영 방식을 옆 지역으로 복제한다.",
    )
    add_table(
        doc,
        ["기간", "목표", "핵심 실행", "통과 기준"],
        [
            ("0~30일", "공급 확보", "이태원 30개 정보, 핵심 제휴 10개, QR 파일럿 3개", "주 2회 이상 정보 갱신"),
            ("31~60일", "첫 수요", "Reels/TikTok, 호스텔·대학·커뮤니티, 현장 QR", "월 활성 1,000명"),
            ("61~90일", "반복 사용", "포인트, 추천 Route, 친구 초대, 주간 이벤트", "30일 재방문 20%"),
            ("4~6개월", "지역 확장", "홍대 복제, 매장 대시보드, 유료 파일럿", "유료 파트너 10개"),
            ("7~12개월", "거래 확장", "티켓·VIP 예약·브랜드 캠페인", "월 거래와 반복 매출 검증"),
        ],
        [1200, 1700, 4100, 2360],
        header_fill=LIGHT,
        first_col_bold=True,
        font_size=9.0,
    )

    add_heading(doc, "7.2 사용자 획득 채널", 2)
    channels = [
        ("숏폼 콘텐츠", "오늘 밤 홍대 3곳, 외국인 입장 가능한 곳, 장르별 추천 등 즉시 저장되는 정보"),
        ("현장 QR", "입구·테이블·화장실·계산대 스티커로 앱 설치와 체크인을 연결"),
        ("크리에이터", "대형 인플루언서보다 지역별 마이크로 크리에이터와 성과형 협업"),
        ("관광 채널", "호스텔·호텔·투어·유학원·대학 국제처에 다국어 QR 가이드 배포"),
        ("검색", "지역+클럽, foreigner-friendly, tonight, guestlist 키워드의 웹 콘텐츠"),
        ("커뮤니티", "외국인 커뮤니티, 대학 동아리, 음악 장르 커뮤니티, DJ 팬덤"),
        ("추천", "친구가 첫 체크인하면 양쪽 모두 포인트를 받는 이중 보상"),
    ]
    for title, detail in channels:
        add_mixed_paragraph(doc, [(f"{title}: ", True), (detail, False)], after=5)

    add_heading(doc, "7.3 콘텐츠 기획", 2)
    for item in [
        "Tonight in Seoul: 오늘 실제로 갈 수 있는 5곳",
        "Foreigner Friendly Verified: 여권·나이·복장·입장 조건이 확인된 장소",
        "오늘의 추천 Route: 21:00 바 → 23:30 클럽 → 03:00 라운지/해장",
        "Genre Night: Hip-hop, House/Techno, K-pop, Live Band, Latin",
        "KR NIGHT Check: 직원에게 직접 확인한 당일 가격·라인업·마감 정보",
        "Neighborhood Battle: 홍대 vs 이태원, 사용자 투표와 저장",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7.4 바이럴 루프", 2)
    add_numbered(doc, "친구 초대 → 친구의 첫 QR 체크인 → 두 사람 모두 포인트", num_id_value=901)
    add_numbered(doc, "추천 Route 공유 → 친구가 저장·참여 → 그룹 체크인 보너스", num_id_value=901)
    add_numbered(doc, "Lounge에서 친구 전환 → 다음 주 친구 상태·추천으로 재방문", num_id_value=901)
    add_numbered(doc, "매장 이벤트 배너 → 현장 콘텐츠 생성 → 매장과 사용자가 재공유", num_id_value=901)

    add_heading(doc, "7.5 초기 월 마케팅 예산 가설", 2)
    add_table(
        doc,
        ["항목", "린 운영", "기본 운영", "비고"],
        [
            ("콘텐츠 제작", "50만원", "150만원", "촬영·편집·번역"),
            ("마이크로 크리에이터", "50만원", "200만원", "성과형 코드 병행"),
            ("현장물·이벤트", "30만원", "100만원", "QR 스티커·작은 파티"),
            ("유료 광고", "50만원", "200만원", "지역·언어·주말 타겟"),
            ("커뮤니티·앰배서더", "20만원", "100만원", "포인트·활동비"),
            ("합계", "월 200만원", "월 750만원", "CAC와 체크인 기준으로 조정"),
        ],
        [2400, 1800, 1800, 3360],
        header_fill=PALE_PINK,
        first_col_bold=True,
        font_size=9.3,
    )


def section_operations_tech_safety(doc: Document):
    add_kicker(doc, "08 · OPERATIONS, TECH & TRUST", color=CYAN)
    add_heading(doc, "운영, 기술, 안전 및 법적 과제", 1)

    add_heading(doc, "8.1 현실적인 운영 구조", 2)
    add_paragraph(
        doc,
        "초기에는 창업자 1명과 비개발 운영 인력 1명으로 파일럿이 가능하다. 단, 제품 개발·보안·법률·디자인은 "
        "필요할 때 외부 전문가를 쓰고, 30개 이상 제휴와 실시간 고객 문의가 생기면 운영 인력을 늘린다.",
    )
    add_table(
        doc,
        ["역할", "초기 담당", "주요 업무"],
        [
            ("대표/PM", "창업자", "제품 우선순위, 영업, 핵심 파트너, 데이터 리뷰"),
            ("현장 운영", "파트타임/첫 직원", "매장 정보 검수, 제휴 연락, QR 설치, 주말 모니터링"),
            ("개발", "창업자+외주/파트너", "백엔드, 인증, QR, 채팅, 앱 배포, 보안"),
            ("콘텐츠", "운영 겸임+프리랜서", "촬영, 숏폼, 번역, Route 큐레이션"),
            ("법률·회계", "전문가", "약관, 개인정보, 위치정보, 결제·세무 검토"),
        ],
        [1700, 2300, 5360],
        header_fill=LIGHT,
        first_col_bold=True,
        font_size=9.4,
    )

    add_heading(doc, "8.2 현재 MVP와 실제 서비스 사이의 간격", 2)
    add_paragraph(
        doc,
        "현재 프로젝트에는 모바일 웹 UI, 4개 언어 전환, 장소 카드, 친구·라운지·게임 화면, 로컬 계정, "
        "QR 미리보기, 포인트·체크인 흐름, Capacitor 기반 iOS 셸, Railway 배포 구성이 준비돼 있다. "
        "그러나 계정·QR·체크인·채팅 데이터는 아직 실제 서버와 운영자 도구로 연결되지 않은 프로토타입이다.",
    )
    add_callout(
        doc,
        "실서비스 전환의 기준",
        "화면이 앱처럼 보이는 것과 실제 서비스가 되는 것은 다르다. 실제 서비스가 되려면 서버 인증, "
        "데이터베이스, 실시간 채팅, 직원 스캐너, 관리자 CMS, 로그·모니터링, 약관·개인정보 체계가 필요하다.",
        fill=PALE_BLUE,
        accent=CYAN,
    )

    add_heading(doc, "8.3 권장 기술 구조", 2)
    tech_rows = [
        ("클라이언트", "현재 Capacitor iOS 셸 유지, Android 추가; 성장 시 React Native/네이티브 재검토"),
        ("백엔드", "Supabase 또는 Firebase로 인증·DB·실시간·스토리지·서버 함수 구성"),
        ("로그인", "Apple, Google, Kakao, NAVER, LINE 우선; WeChat은 중국 사업 필요성 확인 후"),
        ("지도", "한국 사용자는 Kakao/Naver, 외국인은 Google 지도 연결을 병행"),
        ("QR", "서버 서명된 단기 토큰, 스캔 권한이 있는 직원 계정, 중복·재생 방지"),
        ("관리자", "매장 정보, 이벤트, 쿠폰, 신고, 체크인, 포인트를 관리하는 웹 콘솔"),
        ("분석", "노출→상세→길찾기→체크인→재방문 퍼널과 매장별 대시보드"),
    ]
    add_table(
        doc,
        ["구성", "권장안"],
        tech_rows,
        [1700, 7660],
        header_fill=PALE_PINK,
        first_col_bold=True,
        font_size=9.5,
    )

    add_heading(doc, "8.4 나이 확인과 신원", 2)
    add_paragraph(
        doc,
        "초기 전체 가입자에게 여권 KYC를 요구하면 외국인 이탈과 오류가 커질 수 있다. 계정 가입은 낮은 마찰로 "
        "유지하되, 실제 주류 판매·클럽 입장 연령 확인은 매장 직원의 실물 신분증 확인을 기본으로 둔다. "
        "매칭·VIP 우선입장·티켓 거래처럼 위험이나 금전이 커지는 기능에는 선택형 19+ 인증을 단계적으로 도입한다.",
    )
    add_source_note(doc, "주의: 실제 연령 확인 책임과 방식은 업종·기능에 따라 법률 전문가와 매장 계약으로 확정해야 한다.")

    add_heading(doc, "8.5 위치·개인정보 원칙", 2)
    for item in [
        "위치 공유는 기본 OFF이며 사용자가 직접 켜야 한다.",
        "친구 관계는 상호 승인 방식으로 하고, 공개 사용자에게 정확한 위치를 노출하지 않는다.",
        "기본 상태는 좌표가 아닌 매장명 또는 지역 단위로 표시한다.",
        "공유 대상·정확도·종료 시간을 사용자가 선택하고 주기적으로 재동의를 확인한다.",
        "체크아웃·시간 만료·앱 설정 변경 시 즉시 위치 상태와 라운지 권한을 종료한다.",
        "수집 목적, 보유 기간, 제3자 제공, 삭제 방법을 별도 동의와 정책에 명시한다.",
    ]:
        add_bullet(doc, item)
    add_paragraph(
        doc,
        "Snapchat도 위치 공유를 기본 OFF로 두고 상호 친구 중 선택한 사람에게만 공개하며, 정기적으로 "
        "공유 의사를 재확인하는 안전 원칙을 사용한다.[6] 국내에서는 개인위치정보를 활용하는 서비스가 "
        "위치정보법상 신고·약관·보호조치 대상이 될 수 있다. 소상공인·1인 창조기업 예외 조항이 존재하더라도 "
        "서비스 구조에 따른 적용 여부를 출시 전 확인해야 한다.[8]",
    )
    add_source_note(doc, "법률·규제 관련 내용은 제품 기획상의 체크리스트이며 법률 자문이 아니다.[6][8]")

    add_heading(doc, "8.6 커뮤니티 안전", 2)
    add_table(
        doc,
        ["위험", "예방 설계", "운영 대응"],
        [
            ("괴롭힘·성희롱", "상호 친구, DM 제한, 금칙어, 신고 버튼", "신속 차단, 계정·기기 제재, 매장 협조"),
            ("스토킹·위치 노출", "위치 OFF 기본, 장소 단위, 자동 만료", "위치 기록 최소화, 긴급 숨김"),
            ("미성년자", "19+ 안내, 매장 신분증 확인, 기능별 인증", "의심 계정 정지, 매장 정책 연동"),
            ("QR 부정 사용", "회전 토큰, 직원 권한, 재생 방지", "로그 검토, 포인트 회수, 계정 정지"),
            ("과음·안전사고", "책임 음주 문구, 긴급 연락, 귀가 정보", "매장·보안·응급기관 절차"),
        ],
        [1700, 3800, 3860],
        header_fill=LIGHT,
        first_col_bold=True,
        font_size=9.1,
    )


def section_finance_kpi_roadmap(doc: Document):
    add_kicker(doc, "09 · NUMBERS & ROADMAP", color=PINK)
    add_heading(doc, "재무 가설, KPI, 실행 로드맵", 1)

    add_heading(doc, "9.1 월 운영비 가설", 2)
    add_table(
        doc,
        ["항목", "초기 범위/월", "설명"],
        [
            ("서버·지도·메시징", "20만~150만원", "사용량, 지도 API, 이미지·알림·실시간 채팅에 따라 증가"),
            ("운영 인력", "250만~600만원", "1명 또는 파트타임 조합; 창업자 급여 제외"),
            ("콘텐츠·번역", "50만~250만원", "촬영·편집·4개 언어 검수"),
            ("마케팅", "200만~750만원", "린 운영~기본 운영 가설"),
            ("법률·회계·도구", "50만~300만원", "초기 약관·계약 검토 시 일시 증가"),
            ("합계", "월 570만~2,050만원", "인건비와 광고 강도에 따라 큰 차이"),
        ],
        [2600, 2100, 4660],
        header_fill=PALE_BLUE,
        first_col_bold=True,
        font_size=9.4,
    )

    add_heading(doc, "9.2 12개월 차 월 매출 시나리오", 2)
    add_paragraph(
        doc,
        "아래는 월 반복 매출과 캠페인·거래 매출을 합친 월 매출 런레이트 가설이다. 시장 크기나 성공을 "
        "보장하지 않으며, 유료 전환율과 실제 체크인 데이터를 얻기 전에는 투자 판단 근거로 사용할 수 없다.",
    )
    add_table(
        doc,
        ["시나리오", "매장 구독", "광고·캠페인", "성과·거래", "사용자 VIP", "월 합계"],
        [
            ("보수", "144만원", "80만원", "70만원", "138만원", "432만원"),
            ("기준", "720만원", "300만원", "400만원", "552만원", "1,972만원"),
            ("공격", "2,200만원", "1,000만원", "1,500만원", "1,725만원", "6,425만원"),
        ],
        [1200, 1620, 1620, 1620, 1620, 1680],
        header_fill=PALE_PINK,
        first_col_bold=True,
        font_size=9.1,
    )
    add_source_note(
        doc,
        "산식 예시: 기준 시나리오 = 유료 매장 40곳×월 평균 18만원 + 캠페인 300만원 + 거래 400만원 "
        "+ VIP 800명×6,900원. 세금·결제수수료·환불·매장 정산 전 매출.",
    )

    add_heading(doc, "9.3 핵심 KPI", 2)
    add_table(
        doc,
        ["영역", "North Star / 핵심 지표", "초기 목표 가설"],
        [
            ("핵심 가치", "월간 유효 QR 체크인 수", "90일 내 500건"),
            ("공급", "활성 제휴 매장 수·정보 갱신율", "30곳 / 주간 80%"),
            ("획득", "설치 또는 가입당 비용(CAC)", "채널별 추적 후 5천원 이하 검증"),
            ("활성", "가입 후 7일 내 장소 상세·길찾기·체크인", "30% 이상"),
            ("전환", "상세 조회→길찾기→체크인", "단계별 이탈 원인 측정"),
            ("유지", "30일 재방문율", "20% 이상"),
            ("수익", "유료 매장 전환율·매장당 월 매출", "파일럿 종료 후 20% / 15만원+"),
            ("품질", "신고율·정보 오류율·QR 부정률", "각 1% 미만 지향"),
        ],
        [1500, 4660, 3200],
        header_fill=LIGHT,
        first_col_bold=True,
        font_size=9.1,
    )

    add_heading(doc, "9.4 8주 실행 로드맵", 2)
    add_table(
        doc,
        ["주차", "제품", "현장·영업", "성공 기준"],
        [
            ("1~2주", "백엔드·계정·장소 DB 설계", "이태원 30곳 정보 검수, 10곳 미팅", "정보 구조와 제휴 제안 확정"),
            ("3~4주", "실제 로그인, 관리자 CMS, QR 체크인", "파일럿 3곳 직원 교육", "현장 스캔 95% 성공"),
            ("5~6주", "포인트, Lounge, 신고·차단", "QR 설치, 첫 캠페인, 콘텐츠 20개", "유효 체크인 100건"),
            ("7주", "지도·친구·위치 동의, 분석", "사용자 인터뷰 20명, 매장 리뷰", "핵심 오류·이탈 원인 정리"),
            ("8주", "안정화·앱 배포 준비", "유료 파일럿 제안, 홍대 후보 수집", "재사용 의향과 유료 의향 확인"),
        ],
        [1100, 2750, 3150, 2360],
        header_fill=PALE_BLUE,
        first_col_bold=True,
        font_size=9.0,
    )

    add_heading(doc, "9.5 Go / No-Go 판단 기준", 2)
    for item in [
        "Go: 3개 이상 매장에서 직원이 QR 운영을 문제없이 반복한다.",
        "Go: 사용자 100명 중 의미 있는 비율이 다시 앱을 열거나 친구에게 공유한다.",
        "Go: 최소 3개 매장이 무료 기간 후 비용 지불 또는 성과형 계약 의사를 보인다.",
        "No-Go/수정: QR 체크인이 입장 흐름을 방해하거나 매장 직원이 사용을 거부한다.",
        "No-Go/수정: 사용자는 정보만 보고 체크인·포인트·라운지에는 반응하지 않는다.",
        "No-Go/수정: 안전·신고 운영 비용이 감당 가능한 수준을 넘는다.",
    ]:
        add_bullet(doc, item)


def section_partner_pitch_conclusion(doc: Document):
    add_kicker(doc, "10 · PARTNERSHIP & CONCLUSION", color=CYAN)
    add_heading(doc, "제휴 제안과 결론", 1)

    add_heading(doc, "10.1 매장에 제안할 한 문장", 2)
    add_callout(
        doc,
        "제휴 메시지",
        "KR NIGHT가 한국인과 외국인에게 매장의 오늘 이벤트와 입장 정보를 정확히 보여주고, "
        "QR 체크인으로 실제 방문과 재방문 성과까지 확인해드리겠습니다.",
        fill=PALE_PINK,
        accent=PINK,
    )

    add_heading(doc, "10.2 초기 제휴 패키지", 2)
    for item in [
        "무료 매장 페이지와 한국어·영어 기본 정보 정리",
        "당일 이벤트·라인업·입장료·드레스코드 업데이트",
        "KR NIGHT 전용 쿠폰 또는 웰컴 혜택 1개",
        "입구 QR 체크인 설치와 직원 사용 안내",
        "월간 노출·상세 조회·길찾기·체크인·쿠폰 사용 리포트",
        "숏폼 콘텐츠 또는 추천 Route 1회 테스트",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "10.3 매장이 부담해야 할 것", 2)
    for item in [
        "정확한 영업·입장·가격 정보 제공과 변경 시 빠른 알림",
        "직원의 QR 스캔 및 실물 신분증 확인 협조",
        "쿠폰 혜택 이행과 분쟁 시 고객 응대",
        "안전사고·괴롭힘 신고 시 KR NIGHT와의 공동 대응",
        "성과 리포트를 위한 최소 4주 파일럿 참여",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "10.4 KR NIGHT가 지켜야 할 원칙", 2)
    principles = [
        ("정확성", "‘핫하다’는 과장보다 오늘 실제로 방문 가능한 정보를 우선한다."),
        ("측정", "노출 수가 아니라 유효 체크인과 재방문으로 가치를 증명한다."),
        ("동의", "위치·친구·소셜은 모두 선택형이며 언제든 종료할 수 있어야 한다."),
        ("안전", "성장보다 연령·신고·차단·매장 협력 체계를 먼저 갖춘다."),
        ("밀도", "서울 전체의 얕은 목록보다 한 지역의 깊고 최신인 정보를 만든다."),
        ("독립성", "벤치마킹하되 브랜드, 캐릭터, UX, 데이터 구조는 KR NIGHT만의 것으로 만든다."),
    ]
    for title, detail in principles:
        add_mixed_paragraph(doc, [(f"{title}: ", True), (detail, False)], after=5)

    add_heading(doc, "10.5 최종 결론", 2)
    add_paragraph(
        doc,
        "KR NIGHT의 가능성은 ‘클럽을 소개하는 앱’에 있지 않다. 사용자의 밤 선택을 돕고, 실제 방문을 "
        "증명하고, 현장에서 관계와 혜택을 만들고, 그 결과를 매장 매출로 연결하는 구조에 있다. "
        "이 구조가 작동하면 매장은 광고를 사는 것이 아니라 검증된 고객 유입 시스템에 비용을 지불하게 된다.",
    )
    add_paragraph(
        doc,
        "따라서 다음 행동은 기능을 더 많이 그리는 것이 아니라, 이태원 핵심 매장 3곳에서 실제 QR "
        "체크인과 포인트를 운영하고, 사용자의 재사용 의향과 매장의 지불 의향을 확인하는 것이다. "
        "KR NIGHT는 크게 시작할 필요는 없지만, 처음부터 측정 가능한 방식으로 시작해야 한다.",
    )
    add_callout(
        doc,
        "다음 결정",
        "실제 백엔드·직원 스캐너·관리자 CMS를 연결한 파일럿을 만들고, 이태원 제휴 3곳에서 4주간 "
        "‘정보 → 체크인 → 혜택 → 재방문’ 전체 흐름을 검증한다.",
        fill=PALE_BLUE,
        accent=CYAN,
    )


def section_sources(doc: Document):
    add_kicker(doc, "APPENDIX", color=PINK)
    add_heading(doc, "출처 및 참고자료", 1)
    add_paragraph(
        doc,
        "시장·벤치마킹·안전 원칙은 아래 공개 자료를 참고했다. 웹 자료는 2026년 7월 3일 기준으로 확인했다.",
    )
    sources = [
        (
            "[1] 서울특별시, Seoul Welcomes 1.56 Million Visitors in April, Generating KRW 1.15 Trillion in Tourist Spending",
            "https://english.seoul.go.kr/seoul-welcomes-1-56-million-visitors-in-april-generating-krw-1-15-trillion-in-tourist-spending/",
        ),
        ("[2] JP NIGHT 공식 사이트", "https://www.jp-night.com/"),
        ("[3] 덤즈 Dumbs 공식 사이트", "https://www.dumbs.app/"),
        ("[4] CatchTable Japan 공식 사이트", "https://www.catchtable.jp/"),
        ("[5] Resident Advisor 공식 사이트", "https://ra.co/"),
        (
            "[6] Snap Inc., Looking Out for Friends on the Snap Map",
            "https://values.snap.com/news/looking-out-for-friends-on-the-snap-map",
        ),
        ("[7] NightFlow Seoul 공식 사이트", "https://nightflow.kr/"),
        (
            "[8] 국가법령정보센터, 위치정보의 보호 및 이용 등에 관한 법률 제9조의2",
            "https://www.law.go.kr/LSW/LsiJoLinkP.do?docType=JO&joNo=000900000&languageType=KO&lsNm=%EC%9C%84%EC%B9%98%EC%A0%95%EB%B3%B4%EC%9D%98+%EB%B3%B4%ED%98%B8+%EB%B0%8F+%EC%9D%B4%EC%9A%A9+%EB%93%B1%EC%97%90+%EA%B4%80%ED%95%9C+%EB%B2%95%EB%A5%A0&paras=1",
        ),
    ]
    source_table = doc.add_table(rows=0, cols=2)
    set_table_geometry(source_table, [8200, 1160])
    set_table_borders(source_table, color=GRID, size=4)
    for title, url in sources:
        row = source_table.add_row()
        left, right = row.cells
        left_p = left.paragraphs[0]
        left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_p.paragraph_format.space_before = Pt(1)
        left_p.paragraph_format.space_after = Pt(1)
        left_p.paragraph_format.line_spacing = 1.05
        set_run_font(left_p.add_run(title), size=8.5, color=INK, bold=True)
        right_p = right.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        right_p.paragraph_format.space_before = Pt(1)
        right_p.paragraph_format.space_after = Pt(1)
        right_p.paragraph_format.line_spacing = 1.05
        add_hyperlink(right_p, "열기", url)
    source_spacer = doc.add_paragraph()
    source_spacer.paragraph_format.space_after = Pt(2)

    add_heading(doc, "문서 사용 시 주의", 2)
    for item in [
        "본 문서는 사업 전략 초안이며 투자 권유서, 법률 의견서, 확정 재무예측이 아니다.",
        "가격·매출·KPI 목표는 사용자·매장 인터뷰와 실제 파일럿 데이터에 따라 수정해야 한다.",
        "위치정보, 개인정보, 연령 확인, 주류·티켓·결제 관련 사항은 출시 전 전문 검토가 필요하다.",
        "경쟁 서비스의 기능·정책은 변경될 수 있으므로 제휴·투자 설명 전 최신 정보를 다시 확인한다.",
    ]:
        add_bullet(doc, item)


def audit_tables(doc: Document):
    for index, table in enumerate(doc.tables, start=1):
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
        if tbl_w is None or tbl_w.get(qn("w:w")) != str(CONTENT_WIDTH_DXA):
            raise AssertionError(f"Table {index}: bad tblW")
        if tbl_ind is None or tbl_ind.get(qn("w:w")) != str(TABLE_INDENT_DXA):
            raise AssertionError(f"Table {index}: bad tblInd")
        grid_widths = [int(node.get(qn("w:w"))) for node in table._tbl.tblGrid]
        if sum(grid_widths) != CONTENT_WIDTH_DXA:
            raise AssertionError(f"Table {index}: bad grid width sum {sum(grid_widths)}")
        for row in table.rows:
            for cell_index, cell in enumerate(row.cells):
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                if tc_w is None or int(tc_w.get(qn("w:w"))) != grid_widths[cell_index]:
                    raise AssertionError(f"Table {index}: bad cell width")


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    QA_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_document_map(doc)
    section_executive_summary(doc)
    section_problem_intent(doc)
    section_targets(doc)
    section_product(doc)
    section_market_benchmark(doc)
    section_business_model(doc)
    section_marketing(doc)
    section_operations_tech_safety(doc)
    section_finance_kpi_roadmap(doc)
    section_partner_pitch_conclusion(doc)
    section_sources(doc)
    audit_tables(doc)

    core = doc.core_properties
    core.title = "KR NIGHT 사업기획서"
    core.subject = "한국 나이트라이프 발견·체크인·소셜 플랫폼 사업기획"
    core.author = "KR NIGHT"
    core.keywords = "KR NIGHT, nightlife, QR check-in, Seoul, business plan"
    core.comments = "Version 1.0, 2026-07-03"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
